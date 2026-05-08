"""
文件解析服务 - 解析上传的基因列表文件
支持格式：txt, csv, xlsx, docx
包含 Excel 陷阱自动清洗
"""

import os
import re

# 已知的 Excel 日期转换陷阱
EXCEL_DATE_GENES = {
    "sept1": "SEPT1", "sept2": "SEPT2", "sept3": "SEPT3", "sept4": "SEPT4",
    "sept5": "SEPT5", "sept6": "SEPT6", "sept7": "SEPT7", "sept8": "SEPT8",
    "sept9": "SEPT9", "sept10": "SEPT10", "sept11": "SEPT11", "sept12": "SEPT12",
    "sept14": "SEPT14", "1-sep": "SEPT1", "2-sep": "SEPT2", "3-sep": "SEPT3",
    "4-sep": "SEPT4", "5-sep": "SEPT5", "6-sep": "SEPT6", "7-sep": "SEPT7",
    "8-sep": "SEPT8", "9-sep": "SEPT9", "10-sep": "SEPT10", "11-sep": "SEPT11",
    "12-sep": "SEPT12", "14-sep": "SEPT14",
    "1-mar": "MARCH1", "2-mar": "MARCH2", "3-mar": "MARCH3", "4-mar": "MARCH4",
    "5-mar": "MARCH5", "6-mar": "MARCH6", "7-mar": "MARCH7", "8-mar": "MARCH8",
    "9-mar": "MARCH9", "1-may": "MAY1", "2-may": "MAY2", "3-may": "MAY3",
    "1-dec": "DEC1", "2-dec": "DEC2", "3-dec": "DEC3", "4-dec": "DEC4",
    "5-dec": "DEC5", "6-dec": "DEC6", "7-dec": "DEC7", "8-dec": "DEC8",
    "9-dec": "DEC9", "10-dec": "DEC10", "11-dec": "DEC11", "12-dec": "DEC12",
    "1-jan": "JAN1", "2-jan": "JAN2", "3-jan": "JAN3", "4-jan": "JAN4",
    "5-jan": "JAN5", "6-jan": "JAN6", "7-jan": "JAN7", "8-jan": "JAN8",
    "1-jun": "JUN1", "2-jun": "JUN2", "3-jun": "JUN3", "4-jun": "JUN4",
    "5-jun": "JUN5", "6-jun": "JUN6", "7-jun": "JUN7", "8-jun": "JUN8",
    "1-apr": "APR1", "2-apr": "APR2", "3-apr": "APR3", "4-apr": "APR4",
    "5-apr": "APR5", "6-apr": "APR6", "7-apr": "APR7",
    "1-feb": "FEB1", "2-feb": "FEB2", "3-feb": "FEB3", "4-feb": "FEB4",
    "5-feb": "FEB5", "6-feb": "FEB6", "7-feb": "FEB7",
    "1-jul": "JUL1", "2-jul": "JUL2", "3-jul": "JUL3", "4-jul": "JUL4",
    "5-jul": "JUL5", "6-jul": "JUL6",
    "1-aug": "AUG1", "2-aug": "AUG2", "3-aug": "AUG3", "4-aug": "AUG4",
    "5-aug": "AUG5", "6-aug": "AUG6", "7-aug": "AUG7", "8-aug": "AUG8",
    "1-oct": "OCT1", "2-oct": "OCT2", "3-oct": "OCT3", "4-oct": "OCT4",
    "5-oct": "OCT5", "6-oct": "OCT6", "7-oct": "OCT7", "8-oct": "OCT8",
    "9-oct": "OCT9", "10-oct": "OCT10", "11-oct": "OCT11",
    "1-nov": "NOV1", "2-nov": "NOV2", "3-nov": "NOV3", "4-nov": "NOV4",
    "5-nov": "NOV5", "6-nov": "NOV6", "7-nov": "NOV7",
    "1-pip": "PIP1", "2-pip": "PIP2", "3-pip": "PIP3",
    "sep-04": "SEPT4", "sep-05": "SEPT5", "sep-06": "SEPT6", "sep-07": "SEPT7",
    "sep-08": "SEPT8", "sep-09": "SEPT9", "sep-10": "SEPT10", "sep-11": "SEPT11",
    "sep-12": "SEPT12", "sep-14": "SEPT14",
    "mar-01": "MARCH1", "mar-02": "MARCH2", "mar-03": "MARCH3",
    "mar-04": "MARCH4", "mar-05": "MARCH5", "mar-06": "MARCH6",
    "mar-07": "MARCH7", "mar-08": "MARCH8",
}

# 日期格式正则
_DATE_PATTERN = re.compile(
    r"^(\d{1,2})[-/](\w{3,9})[-/]?(\d{2,4})?$|^(\w{3,9})[-/](\d{1,2})[-/]?(\d{2,4})?$",
    re.IGNORECASE
)


def clean_gene_name(raw: str) -> dict:
    """
    清洗单个基因名，检测 Excel 陷阱和格式问题

    Returns:
        {
            "original": str,      # 原始值
            "cleaned": str,        # 清洗后的值
            "is_clean": bool,      # 是否无需清洗
            "issue": str | None,   # 问题描述
            "fix": str | None,     # 修复建议
        }
    """
    original = raw.strip()
    if not original:
        return {"original": original, "cleaned": "", "is_clean": True, "issue": None, "fix": None}

    # 去除引号和多余空白
    cleaned = original.strip('"').strip("'").strip()

    # 检查 Excel 日期陷阱
    lower = cleaned.lower()
    if lower in EXCEL_DATE_GENES:
        return {
            "original": original,
            "cleaned": EXCEL_DATE_GENES[lower],
            "is_clean": False,
            "issue": "Excel date conversion detected",
            "fix": f'"{original}" → "{EXCEL_DATE_GENES[lower]}"',
        }

    # 检查日期格式 (e.g., "2024-03-15", "3/15/2024", "15-Mar")
    if _DATE_PATTERN.match(cleaned):
        # 尝试从日期中提取可能的基因名
        parts = re.split(r"[-/]", cleaned)
        for part in parts:
            p = part.strip()
            if p.upper() in EXCEL_DATE_GENES.values():
                return {
                    "original": original,
                    "cleaned": p.upper(),
                    "is_clean": False,
                    "issue": "Date format detected",
                    "fix": f'"{original}" → "{p.upper()}"',
                }

        return {
            "original": original,
            "cleaned": cleaned,
            "is_clean": False,
            "issue": "Possible date format (not a gene)",
            "fix": "Please verify this entry manually",
        }

    # 检查纯数字
    if cleaned.isdigit():
        return {
            "original": original,
            "cleaned": cleaned,
            "is_clean": False,
            "issue": "Numeric value (possibly Entrez Gene ID)",
            "fix": "Use Universal ID Bridge to resolve",
        }

    # 检查乱码/特殊字符
    if re.search(r'[^\x20-\x7E\u00C0-\u024F]', cleaned):
        return {
            "original": original,
            "cleaned": cleaned,
            "is_clean": False,
            "issue": "Contains non-ASCII characters",
            "fix": "Review and remove special characters",
        }

    return {"original": original, "cleaned": cleaned, "is_clean": True, "issue": None, "fix": None}


def parse_gene_file(filepath: str) -> dict:
    """
    解析上传的基因列表文件，包含清洗预览

    Returns:
        {
            "total_lines": int,
            "genes": list[str],          # 清洗后的基因名
            "gene_count": int,
            "filename": str,
            "clean_report": list[dict],  # 清洗报告
            "issues_count": int,         # 有问题的行数
        }
    """
    ext = os.path.splitext(filepath)[1].lower()
    filename = os.path.basename(filepath)

    if ext in (".txt", ".csv"):
        raw_genes = _parse_text_csv_raw(filepath, ext)
    elif ext == ".xlsx":
        raw_genes = _parse_excel_raw(filepath)
    elif ext == ".docx":
        raw_genes = _parse_word_raw(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported: txt, csv, xlsx, docx")

    # 清洗每个基因名
    genes = []
    clean_report = []
    for raw in raw_genes:
        result = clean_gene_name(raw)
        clean_report.append(result)
        if result["cleaned"]:
            genes.append(result["cleaned"])

    issues_count = sum(1 for r in clean_report if not r["is_clean"])

    return {
        "total_lines": len(raw_genes),
        "genes": genes,
        "gene_count": len(genes),
        "filename": filename,
        "clean_report": clean_report,
        "issues_count": issues_count,
    }


def _parse_text_csv_raw(filepath: str, ext: str) -> list[str]:
    """解析 txt/csv 返回原始行"""
    raw = []
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            stripped = line.strip()
            if not stripped:
                continue
            if ext == ".csv":
                parts = stripped.split(",")
                stripped = parts[0].strip()
            stripped = stripped.strip('"').strip("'").strip()
            if stripped:
                raw.append(stripped)
    return raw


def _parse_excel_raw(filepath: str) -> list[str]:
    """解析 Excel 返回原始值"""
    from openpyxl import load_workbook
    wb = load_workbook(filepath, read_only=True)
    ws = wb.active
    raw = []
    for row in ws.iter_rows(values_only=True):
        if row and row[0] is not None:
            val = str(row[0]).strip()
            if val:
                raw.append(val)
    wb.close()
    return raw


def _parse_word_raw(filepath: str) -> list[str]:
    """解析 Word 返回原始值"""
    from docx import Document
    doc = Document(filepath)
    raw = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            raw.append(text)
    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text:
                raw.append(row.cells[0].text.strip())
    return raw
